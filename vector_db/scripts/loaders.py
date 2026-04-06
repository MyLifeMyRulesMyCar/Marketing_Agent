"""
scripts/loaders.py — Extract text from PDF, Excel, Word, TXT, MD, CSV, and URLs.
Returns a list of {"text": str, "source": str, "metadata": dict}
"""

import os
import csv
import requests
from pathlib import Path
from typing import Optional


# ── PDF ──────────────────────────────────────────────────────

def load_pdf(path: str) -> list[dict]:
    try:
        import pdfplumber
        chunks = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if text.strip():
                    chunks.append({
                        "text": text.strip(),
                        "source": path,
                        "metadata": {"type": "pdf", "page": i + 1, "file": os.path.basename(path)}
                    })
        return chunks
    except Exception as e:
        print(f"  ✗ PDF load failed ({path}): {e}")
        return []


# ── Excel ─────────────────────────────────────────────────────

def load_excel(path: str) -> list[dict]:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True)
        chunks = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(c) for c in row if c is not None)
                if row_text.strip():
                    rows.append(row_text)
            if rows:
                # Group rows into chunks
                text = "\n".join(rows)
                chunks.append({
                    "text": text,
                    "source": path,
                    "metadata": {"type": "excel", "sheet": sheet_name, "file": os.path.basename(path)}
                })
        return chunks
    except Exception as e:
        print(f"  ✗ Excel load failed ({path}): {e}")
        return []


# ── Word (docx) ───────────────────────────────────────────────

def load_docx(path: str) -> list[dict]:
    try:
        from docx import Document
        doc = Document(path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        text = "\n".join(paragraphs)
        if not text.strip():
            return []
        return [{
            "text": text,
            "source": path,
            "metadata": {"type": "docx", "file": os.path.basename(path)}
        }]
    except Exception as e:
        print(f"  ✗ DOCX load failed ({path}): {e}")
        return []


# ── Plain text / Markdown ─────────────────────────────────────

def load_text(path: str) -> list[dict]:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read().strip()
        if not text:
            return []
        ext = Path(path).suffix.lower()
        return [{
            "text": text,
            "source": path,
            "metadata": {"type": ext.lstrip("."), "file": os.path.basename(path)}
        }]
    except Exception as e:
        print(f"  ✗ Text load failed ({path}): {e}")
        return []


# ── CSV ───────────────────────────────────────────────────────

def load_csv(path: str) -> list[dict]:
    try:
        rows = []
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            for row in reader:
                line = " | ".join(c.strip() for c in row if c.strip())
                if line:
                    rows.append(line)
        text = "\n".join(rows)
        if not text:
            return []
        return [{
            "text": text,
            "source": path,
            "metadata": {"type": "csv", "file": os.path.basename(path)}
        }]
    except Exception as e:
        print(f"  ✗ CSV load failed ({path}): {e}")
        return []


# ── URL (web page) ────────────────────────────────────────────

def load_url(url: str) -> list[dict]:
    try:
        from bs4 import BeautifulSoup
        headers = {"User-Agent": "Mozilla/5.0 (compatible; VectorDB-Builder/1.0)"}
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()

        soup = BeautifulSoup(resp.text, "html.parser")

        # Remove nav, footer, scripts
        for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
            tag.decompose()

        text = soup.get_text(separator="\n")
        # Clean up whitespace
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        text = "\n".join(lines)

        if not text:
            return []

        return [{
            "text": text,
            "source": url,
            "metadata": {"type": "url", "url": url}
        }]
    except Exception as e:
        print(f"  ✗ URL load failed ({url}): {e}")
        return []


# ── Router ────────────────────────────────────────────────────

LOADERS = {
    ".pdf":  load_pdf,
    ".xlsx": load_excel,
    ".xls":  load_excel,
    ".docx": load_docx,
    ".doc":  load_docx,
    ".txt":  load_text,
    ".md":   load_text,
    ".csv":  load_csv,
}

SUPPORTED = set(LOADERS.keys())


def load_file(path: str) -> list[dict]:
    ext = Path(path).suffix.lower()
    loader = LOADERS.get(ext)
    if loader:
        return loader(path)
    else:
        print(f"  ⚠ Unsupported format: {ext} ({path})")
        return []


def load_folder(folder: str) -> list[dict]:
    """Recursively load all supported files from a folder."""
    docs = []
    folder_path = Path(folder)
    if not folder_path.exists():
        print(f"  ⚠ Folder not found: {folder}")
        return []

    for fpath in sorted(folder_path.rglob("*")):
        if fpath.is_file() and fpath.suffix.lower() in SUPPORTED:
            print(f"  📄 {fpath}")
            docs.extend(load_file(str(fpath)))

    return docs